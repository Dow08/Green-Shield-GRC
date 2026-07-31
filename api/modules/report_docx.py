"""report_docx.py — rapport de mission au format Word (.docx), identité visuelle
alignée sur report_html.py.

Raison d'être : jusqu'au 31/07/2026, le `.docx` passait par un gabarit statique
(`api/templates/rapport_iso27001.docx`) généré en `python-docx` nu — police
par défaut, aucune couleur, 7 sections génériques figées depuis des mois
pendant que `report_html.py` montait à 13 (TPRM, AIPD, E3R, DORA jamais
ajoutés). Constat en recette : le Word ne ressemblait à rien de ce que
l'application produit ailleurs, et son titre affichait « Rapport d'audit de
conformité » même sur une mission de conseil — écrit en dur, indépendant du
volet réel.

Ce module reconstruit le document en `python-docx` directement, sans gabarit
intermédiaire : chaque section lit les mêmes champs de la mission que
`report_html.py`, dans le même ordre. Deux points de passage uniques
garantissent qu'ils ne peuvent plus diverger silencieusement :
  * `report_html.CHAPITRES` — la liste des titres de chapitre, importée telle
    quelle ; un chapitre ajouté au HTML apparaît donc aussi au sommaire Word.
  * `report_html.titre_et_meta()` — même titre, même bandeau méta, sur les
    deux formats.

Aucune ressource externe : le logo est décodé depuis `charte.LOGO_BASE64`
(déjà embarqué pour les livrables HTML/Markdown), pas de police tierce à
installer — Word substitue silencieusement une police absente, ce qui
romprait la mise en page sans avertissement.
"""
from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

from . import aipd as aipd_module
from . import charte
from . import controles_techniques
from . import couverture
from . import docx_export
from . import report_html
from . import soa as soa_module
from . import tprm

# L'application sert n'importe quel consultant, pas un seul cabinet : jamais
# de nom d'entreprise ou de personne écrit en dur dans un document généré.
# Repéré le 30/07/2026 (retour utilisateur) : plusieurs endroits retombaient
# silencieusement sur "Dorian"/"DP Cyber Consulting" au lieu des valeurs
# transmises par Réglages, y compris quand elles étaient vides.
_AUDITEUR_DEFAUT = "Consultant"
_CABINET_DEFAUT = "Cabinet non renseigné"

# --- Palette : mêmes valeurs que report_html.py::_FEUILLE --------------------
_VERT = "2EE6A0"
_VERT_CLAIR = "7BF3C8"
_SOMBRE = "04150E"
_ENCRE = "0C2317"
_CORPS = "33403B"
_DOUX = "6B7F78"
_TRAIT = "DFE7E3"
_TEINTE = "F2FBF7"
_ROSE, _ROSE_F = "A3243C", "FDEAED"
_AMBRE, _AMBRE_F = "8A5B00", "FDF3E0"
_CIEL, _CIEL_F = "1D5F88", "E8F4FD"
_OK, _OK_F = "12694A", "E6F8F0"
_NEUTRE, _NEUTRE_F = "62726D", "EEF2F1"

# Palette Word par classe de sévérité (report_html.CLASSE_SEV décide QUEL
# statut appartient à quelle classe ; ce module décide seulement de la
# couleur Word qui l'illustre — la même distinction que `.sev-*` côté CSS.
_SEV_COULEURS = {
    "crit": (_ROSE, _ROSE_F), "elev": (_AMBRE, _AMBRE_F), "moy": (_CIEL, _CIEL_F),
    "faib": (_NEUTRE, _NEUTRE_F), "ok": (_OK, _OK_F),
}


def _rgb(hexa: str) -> RGBColor:
    return RGBColor.from_string(hexa)


def _largeur_utile(doc: Document) -> Emu:
    """Largeur imprimable de la page (largeur — marges gauche et droite).

    Un tableau `python-docx` créé sans largeur de colonne explicite se répartit
    à parts **égales** entre toutes les colonnes dès l'ouverture dans Word —
    quelle que soit la longueur réelle du contenu. Constaté le 31/07/2026 :
    une colonne « G » à deux caractères recevait la même largeur qu'une colonne
    de scénario tenant sur huit lignes. `_table()` a donc besoin de calculer
    des largeurs de colonne explicites, et pour ça, de connaître l'espace
    réellement disponible sur la page.
    """
    section = doc.sections[0]
    return Emu(section.page_width - section.left_margin - section.right_margin)


def _fixer_largeurs(table, largeurs_emu: list[int]) -> None:
    """Impose des largeurs de colonne fixes, respectées par Word à l'ouverture.

    Deux conditions, aucune facultative :
      * `table.autofit = False` — sinon Word recalcule les largeurs à
        l'ouverture du document plutôt que de garder celles qu'on lui donne ;
      * la largeur posée sur **chaque cellule de chaque ligne**, pas seulement
        sur `table.columns[i]` — cette dernière ne met à jour que les lignes
        déjà présentes au moment de l'appel, jamais celles ajoutées ensuite
        par `table.add_row()`.
    """
    table.autofit = False
    table.allow_autofit = False
    for i, largeur in enumerate(largeurs_emu):
        table.columns[i].width = Emu(largeur)
    for row in table.rows:
        for i, largeur in enumerate(largeurs_emu):
            row.cells[i].width = Emu(largeur)


def _largeurs_ponderees(doc: Document, poids: tuple[float, ...]) -> list[int]:
    """Convertit des poids relatifs (ex. `(1, 3, 1)`) en largeurs de colonne
    en EMU, proportionnelles à l'espace réellement disponible sur la page."""
    disponible = int(_largeur_utile(doc))
    somme = sum(poids)
    return [int(disponible * p / somme) for p in poids]


def _shade(element, hexa: str) -> None:
    """Teinte le fond d'une cellule de tableau (w:shd)."""
    pr = element._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexa)
    pr.append(shd)


def _bordure_basse_cellule(cell, hexa: str, taille: int = 4) -> None:
    """Ligne fine sous une cellule — le style de tableau du produit (une seule
    ligne horizontale par ligne, pas de quadrillage complet)."""
    pr = cell._tc.get_or_add_tcPr()
    bordures = OxmlElement("w:tcBorders")
    bas = OxmlElement("w:bottom")
    bas.set(qn("w:val"), "single")
    bas.set(qn("w:sz"), str(taille))
    bas.set(qn("w:space"), "0")
    bas.set(qn("w:color"), hexa)
    bordures.append(bas)
    pr.append(bordures)


def _bordure_basse_paragraphe(paragraphe, hexa: str, taille: int = 16) -> None:
    """Ligne sous un titre de chapitre — équivalent de `border-bottom` CSS."""
    pr = paragraphe._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bas = OxmlElement("w:bottom")
    bas.set(qn("w:val"), "single")
    bas.set(qn("w:sz"), str(taille))
    bas.set(qn("w:space"), "4")
    bas.set(qn("w:color"), hexa)
    pbdr.append(bas)
    pr.append(pbdr)


def _logo_bytes(logo: str = "") -> io.BytesIO:
    return io.BytesIO(charte.logo_bytes(logo))


def _run(paragraphe, texte: str, *, gras: bool = False, italique: bool = False,
         couleur: str | None = None, taille: float | None = None) -> None:
    run = paragraphe.add_run(texte)
    run.bold = gras
    run.italic = italique
    if couleur:
        run.font.color.rgb = _rgb(couleur)
    if taille:
        run.font.size = Pt(taille)


def _t(valeur) -> str:
    texte = "" if valeur is None else str(valeur).strip()
    return texte if texte else "—"


def _libelle_sev(valeur) -> tuple[str, str | None]:
    """Libellé lisible + classe de couleur, ou None si la valeur n'est pas un statut connu."""
    brut = "" if valeur is None else str(valeur).strip()
    if not brut:
        return "—", None
    classe = report_html.CLASSE_SEV.get(brut)
    libelle = docx_export.STATUS_LABELS.get(brut, brut)
    return libelle, classe


# --- Bâtisseurs de contenu ----------------------------------------------------

def _vide(doc: Document, message: str) -> None:
    p = doc.add_paragraph()
    _run(p, message, italique=True, couleur=_DOUX, taille=9)


def _note(doc: Document, texte: str) -> None:
    p = doc.add_paragraph()
    _run(p, texte, couleur=_DOUX, taille=9)


def _encadre(doc: Document, texte: str, *, alerte: bool = False) -> None:
    """Bandeau teinté — note de couverture technique ou alerte AIPD Art. 36."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _shade(cell, _ROSE_F if alerte else _TEINTE)
    p = cell.paragraphs[0]
    _run(p, texte, gras=alerte, couleur=_ROSE if alerte else _ENCRE, taille=9.5)
    _fixer_largeurs(table, [int(_largeur_utile(doc))])
    doc.add_paragraph()


def _table(doc: Document, entetes: tuple[str, ...], lignes: list[tuple], message_vide: str,
           colonnes_num: tuple[int, ...] = (), colonnes_sev: tuple[int, ...] = (),
           largeurs: tuple[float, ...] | None = None) -> None:
    """Tableau tramé — miroir de `report_html._table()`.

    `largeurs` : poids relatifs de chaque colonne (ex. `(0.7, 2.6, 0.55, 0.55, 2.6)`
    pour ID / Scénario / Gravité / Vraisemblance / Mesure). Sans lui, les
    colonnes se répartiraient à parts égales quelle que soit la longueur de
    leur contenu — voir `_largeur_utile`.
    """
    if not lignes:
        _vide(doc, message_vide)
        return

    table = doc.add_table(rows=1, cols=len(entetes))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    poids = largeurs or tuple(1 for _ in entetes)
    largeurs_colonnes = _largeurs_ponderees(doc, poids)

    for i, entete in enumerate(entetes):
        cell = table.rows[0].cells[i]
        _shade(cell, _TEINTE)
        _bordure_basse_cellule(cell, "C9DBD3", taille=6)
        p = cell.paragraphs[0]
        if i in colonnes_num:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(p, entete, gras=True, couleur=_ENCRE, taille=8)

    for ligne in lignes:
        cells = table.add_row().cells
        for i, valeur in enumerate(ligne):
            p = cells[i].paragraphs[0]
            _bordure_basse_cellule(cells[i], _TRAIT, taille=4)
            if i in colonnes_sev:
                libelle, classe = _libelle_sev(valeur)
                if classe:
                    couleur, fond = _SEV_COULEURS[classe]
                    _shade(cells[i], fond)
                    _run(p, libelle, gras=True, couleur=couleur, taille=8.5)
                else:
                    _run(p, libelle, couleur=_CORPS, taille=9)
            else:
                if i in colonnes_num:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _run(p, _t(valeur), couleur=_CORPS, taille=9)

    _fixer_largeurs(table, largeurs_colonnes)
    doc.add_paragraph()


def _sous_titre(doc: Document, texte: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    _run(p, texte, gras=True, couleur=_ENCRE, taille=10.5)


def _chapitre_titre(doc: Document, numero: int, titre: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4) if numero == 1 else Pt(20)
    p.paragraph_format.space_after = Pt(8)
    _bordure_basse_paragraphe(p, _VERT, taille=16)
    _run(p, f"{numero}. ", gras=True, couleur=_VERT, taille=14)
    _run(p, titre, gras=True, couleur=_ENCRE, taille=14)


def _titre_sans_numero(doc: Document, titre: str) -> None:
    """Titre de section hors sommaire — les blocs « SIGNATURES » des livrables
    courts (NDA, PSSI, AIPD) n'ont pas de numéro de chapitre, à l'image de
    leur en-tête `###` non numéroté dans le Markdown d'origine."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    _bordure_basse_paragraphe(p, _VERT, taille=16)
    _run(p, titre, gras=True, couleur=_ENCRE, taille=13)


# --- Un bâtisseur par chapitre, même contenu que report_html.py -------------

def _ch_synthese(doc: Document, state: dict, steps: dict) -> None:
    resume = ((steps.get("restitution") or {}).get("exec_summary") or "").strip()
    if not resume:
        _vide(doc, "Synthèse non rédigée. Elle se saisit en phase 6 et n'est jamais "
                   "produite automatiquement : elle engage le jugement du consultant.")
        return
    for paragraphe in resume.split("\n"):
        if paragraphe.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _run(p, paragraphe.strip(), couleur=_CORPS, taille=11)


def _ch_cadrage(doc: Document, state: dict, steps: dict) -> None:
    socle = state.get("socle") or {}
    q, c, k = (socle.get("qualification") or {}, socle.get("contractualisation") or {},
               socle.get("kickoff") or {})
    champs = (
        ("Déclencheur de la mission", q.get("declencheur")),
        ("Sponsor exécutif", q.get("sponsor_executif")),
        ("Budget vendu", q.get("budget")),
        ("Maturité constatée à l'entrée", q.get("maturite_actuelle")),
        ("Équipe interne mobilisable", q.get("equipe_interne")),
        ("Échéance cible", q.get("echeance_cible")),
        ("Périmètre inclus", c.get("perimetre_inclus")),
        ("Périmètre explicitement exclu", c.get("perimetre_exclu")),
        ("Modalités d'intervention", c.get("modalites")),
        ("Accès au SI consentis", c.get("acces_si")),
        ("Date de réunion de lancement", k.get("date")),
        ("Gouvernance de la mission", k.get("gouvernance")),
    )
    lignes = [(lib, val) for lib, val in champs if str(val or "").strip()]
    for libelle, liste in (("Livrables contractuels", c.get("livrables")),
                           ("Participants au lancement", k.get("participants"))):
        if liste:
            lignes.append((libelle, " · ".join(str(x) for x in liste)))
    _table(doc, ("Élément de cadrage", "Contenu"), lignes,
          "Le cadrage contractuel de la mission n'a pas été renseigné.",
          largeurs=(1, 2.3))

    _sous_titre(doc, "2.1 Entretiens conduits")
    entretiens = socle.get("entretiens") or []
    _table(doc, ("Rôle rencontré", "Date", "Ce qui a été déclaré"),
          [(e.get("role"), e.get("date"), e.get("synthese")) for e in entretiens],
          "Aucun entretien n'a été consigné : les constats de ce rapport ne sont pas "
          "rattachés à une source déclarative identifiée.",
          largeurs=(1.3, 0.9, 3))


def _ch_patrimoine(doc: Document, state: dict, steps: dict, prefixe: str = "3") -> None:
    """Réutilisé tel quel par `build_ebios_docx` (prefixe="1") : mêmes colonnes,
    mêmes données que le chapitre patrimoine du rapport de mission."""
    cadrage = steps.get("cadrage") or {}
    _sous_titre(doc, f"{prefixe}.1 Valeurs métier")
    _table(doc, ("ID", "Valeur métier", "Description", "Données personnelles"),
          [(a.get("id"), a.get("name"), a.get("description"),
            "Oui" if a.get("is_personal_data") else "Non")
           for a in cadrage.get("assets_metier") or []],
          "Aucune valeur métier n'a été cartographiée.",
          largeurs=(0.7, 1.6, 3, 1.3))
    _sous_titre(doc, f"{prefixe}.2 Biens supports")
    _table(doc, ("ID", "Bien support", "Type", "Description", "Responsable"),
          [(a.get("id"), a.get("name"), a.get("type"), a.get("description"), a.get("owner"))
           for a in cadrage.get("assets_support") or []],
          "Aucun bien support n'a été inventorié.",
          largeurs=(0.7, 1.6, 1, 2.6, 1.4))


def _table_obligations_aipd(doc: Document, donnees: dict) -> None:
    """Contenu (alerte + tableau) des obligations de procédure de l'AIPD.

    Réutilisé sans le titre englobant : chapitre à part entière dans le
    document AIPD autonome (§3), simple sous-section dans le rapport de
    mission (§4.3) — les deux appelants posent leur propre en-tête.
    """
    saisies = {o.get("id"): o for o in donnees.get("obligations") or []}
    etat = aipd_module.etat(donnees)
    alerte = aipd_module.alerte_bloquante(donnees)
    if alerte:
        _encadre(doc, alerte, alerte=True)
    lignes = []
    for obligation in aipd_module.OBLIGATIONS:
        if obligation["conditionnelle"] and not etat["art36_requise"]:
            lignes.append((obligation["libelle"], obligation["reference"],
                           "Non applicable (risque résiduel non élevé)", ""))
            continue
        saisie = saisies.get(obligation["id"], {})
        lignes.append((obligation["libelle"], obligation["reference"],
                       "Fait" if saisie.get("satisfait") else "Reste à faire",
                       saisie.get("commentaire")))
    _table(doc, ("Obligation", "Référence", "État", "Commentaire"), lignes,
          "Aucune obligation renseignée.",
          largeurs=(2, 1.2, 1, 2))


def _ch_aipd(doc: Document, state: dict, steps: dict) -> None:
    diagnostic = steps.get("diagnostic") or {}
    _sous_titre(doc, "4.1 Registre des traitements (RGPD Art. 30)")
    _table(doc, ("ID", "Traitement", "Finalité", "Catégories de données", "Conservation"),
          [(r.get("id"), r.get("name"), r.get("purpose"), r.get("data_categories"),
            r.get("retention")) for r in diagnostic.get("rgpd_register") or []],
          "Aucun traitement n'a été inscrit au registre.",
          largeurs=(0.7, 1.6, 2, 1.8, 1.3))

    _sous_titre(doc, "4.1bis Registre des violations de données (RGPD Art. 33-34)")
    _table(doc, ("ID", "Constatée le", "Nature", "CNIL", "Personnes informées"),
          [(v.get("id"), v.get("date_constat"), v.get("nature"),
            f"Notifiée le {v.get('date_notification_cnil')}" if v.get("notifiee_cnil") else "Non notifiée",
            "Oui" if v.get("personnes_informees") else "Non")
           for v in diagnostic.get("violations") or []],
          "Aucune violation de données n'a été constatée sur cette mission.",
          largeurs=(0.7, 1, 1.6, 1.4, 1.1))

    if not diagnostic.get("aipd_required"):
        _vide(doc, "Aucune analyse d'impact n'est requise sur ce périmètre.")
        return

    donnees = diagnostic.get("aipd") or {}
    _sous_titre(doc, "4.2 Analyse d'impact — les quatre volets")
    _table(doc, ("Volet d'analyse", "Contenu"),
          [(lib, donnees.get(cle)) for lib, cle in (
              ("Description systématique du traitement", "treatment_description"),
              ("Nécessité et proportionnalité", "necessity_eval"),
              ("Risques pour les droits et libertés", "risks_eval"),
              ("Mesures d'atténuation", "mitigation_measures"),
          )],
          "Les volets d'analyse n'ont pas été renseignés.",
          largeurs=(1.6, 2.4))

    _sous_titre(doc, "4.3 Obligations organisationnelles")
    _table_obligations_aipd(doc, donnees)


def _ch_risque(doc: Document, state: dict, steps: dict, prefixe: str = "5") -> None:
    """Réutilisé tel quel par `build_ebios_docx` (prefixe="2")."""
    ebios = steps.get("ebios") or {}
    _sous_titre(doc, f"{prefixe}.1 Événements redoutés")
    _table(doc, ("ID", "Événement redouté", "Gravité", "Impacts"),
          [(e.get("id"), e.get("event"), f"{e.get('gravity')}/4", e.get("impact"))
           for e in ebios.get("redoute_events") or []],
          "Aucun événement redouté n'a été caractérisé.", colonnes_num=(2,),
          largeurs=(0.7, 2, 0.8, 2.5))

    _sous_titre(doc, f"{prefixe}.2 Sources de risque")
    _table(doc, ("ID", "Source de risque", "Objectif visé"),
          [(s.get("id"), s.get("name"), s.get("objective")) for s in ebios.get("risk_sources") or []],
          "Aucune source de risque n'a été caractérisée.",
          largeurs=(0.7, 2, 2))

    _sous_titre(doc, f"{prefixe}.3 Scénarios opérationnels")
    _table(doc, ("ID", "Scénario opérationnel", "G", "V", "Mesure d'atténuation"),
          [(s.get("id"), s.get("event"), f"{s.get('gravity')}/4", f"{s.get('likelihood')}/5",
            s.get("mitigation")) for s in ebios.get("operational_scenarios") or []],
          "Aucun scénario opérationnel n'a été construit.", colonnes_num=(2, 3),
          largeurs=(0.7, 2.6, 0.55, 0.55, 2.6))

    _sous_titre(doc, f"{prefixe}.3bis Traitement des risques (propriétaire, résiduel, décision)")
    _table(doc, ("ID", "Propriétaire", "Résiduel (G/V)", "Stratégie", "Statut"),
          [(s.get("id"), s.get("owner"),
            f"{s.get('gravite_residuelle')}/{s.get('vraisemblance_residuelle')}"
            if s.get("gravite_residuelle") is not None and s.get("vraisemblance_residuelle") is not None else None,
            s.get("strategie_traitement"), s.get("statut"))
           for s in ebios.get("operational_scenarios") or []],
          "Aucun scénario opérationnel n'a été construit.",
          largeurs=(0.7, 1.5, 1, 1.1, 1))

    _sous_titre(doc, f"{prefixe}.4 Cas réels versés au dossier")
    _table(doc, ("Cas réel", "Enseignement retenu pour ce client"),
          [(c.get("case"), c.get("lessons")) for c in ebios.get("case_studies") or []],
          "Aucun cas comparable n'a été versé au dossier.",
          largeurs=(1.6, 3.4))


def _ch_ecosysteme(doc: Document, state: dict, steps: dict) -> None:
    """Restitution des tiers selon le volet (§14.1bis) : ratio ANSSI, ou exigences GRC."""
    tiers = ((steps.get("tprm") or {}).get("tiers")) or []

    if state.get("type") == "grc":
        _note(doc, "Ce volet ne produit aucun score de risque : ni DORA ni NIS2 ne se "
                   "réclament d'EBIOS RM. La conformité se démontre par des preuves.")
        lignes = []
        for t in tiers:
            etat = tprm.conformite(t)
            manquantes = [e["libelle"] for e in (t.get("exigences") or []) if not e.get("satisfait")]
            lignes.append((t.get("name"), f"{etat['satisfaites']}/{etat['total']} ({etat['taux']} %)",
                           "Conforme" if etat["conforme"] else " ; ".join(manquantes) or "—"))
        _table(doc, ("Prestataire", "Exigences satisfaites", "Écarts restants"), lignes,
              "Aucun prestataire n'a été inscrit au registre.", colonnes_num=(1,),
              largeurs=(1.8, 1.5, 2.2))
        return

    _note(doc, "Criticité selon la formule ANSSI : (dépendance × pénétration) / (maturité × confiance).")
    classement = sorted(tiers, key=lambda t: t.get("score", 0), reverse=True)
    _table(doc, ("Tiers", "Criticité", "Ratio", "Dép. / Pén. / Mat. / Conf."),
          [(t.get("name"), t.get("rating"), t.get("score"),
            f"{t.get('dependence')} / {t.get('penetration')} / {t.get('maturity')} / {t.get('trust')}")
           for t in classement],
          "Aucun tiers n'a été évalué.", colonnes_num=(2, 3), colonnes_sev=(1,),
          largeurs=(1.8, 1.1, 0.9, 1.8))


def _ch_resilience(doc: Document, state: dict, steps: dict) -> None:
    bcp = (steps.get("resilience") or {}).get("bcp_strategy") or {}
    e3r = (steps.get("resilience") or {}).get("e3r") or {}
    strategie = (steps.get("resilience") or {}).get("strategie_remediation") or {}
    _sous_titre(doc, "7.1 Cibles de continuité")
    _table(doc, ("Cible de continuité", "Valeur retenue"),
          [(lib, val) for lib, val in (
              ("RTO — durée maximale d'interruption admissible", bcp.get("rto")),
              ("RPO — perte de données maximale admissible", bcp.get("rpo")),
              ("Politique de sauvegarde", bcp.get("backup_policy")),
          ) if str(val or "").strip()],
          "Aucune cible de continuité n'a été définie.",
          largeurs=(2, 2))

    _sous_titre(doc, "7.2 Séquence de remédiation E3R (ANSSI)")
    _table(doc, ("Étape E3R", "Procédure retenue"),
          [(lib, e3r.get(cle)) for lib, cle in (
              ("Endiguement", "endiguement"), ("Éviction", "eviction"),
              ("Éradication", "eradication"), ("Reconstruction", "reconstruction"),
          ) if str(e3r.get(cle) or "").strip()],
          "La séquence de remédiation E3R n'a pas été documentée.",
          largeurs=(1.4, 3.1))

    _sous_titre(doc, "7.3 Volet stratégique — arbitrage Direction")
    _table(doc, ("Critère d'arbitrage", "Position retenue"),
          [(lib, strategie.get(cle)) for lib, cle in (
              ("Urgence de redémarrage", "urgence_redemarrage"),
              ("Coûts et risques d'un redémarrage précipité", "couts_risques_redemarrage"),
              ("Décision retenue et autorité", "decision_direction"),
          ) if str(strategie.get(cle) or "").strip()],
          "Le volet stratégique (arbitrage Direction) n'a pas été documenté.",
          largeurs=(1.8, 2.7))


def _ch_evaluation(doc: Document, state: dict, steps: dict) -> None:
    controles = (steps.get("evaluation") or {}).get("manual_controls") or []
    _table(doc, ("ID", "Exigence organisationnelle", "Statut", "Constat et preuve"),
          [(c.get("id"), c.get("title"), c.get("status"), c.get("notes")) for c in controles],
          "Aucune check-list de conformité n'est rattachée à cette mission : "
          "l'évaluation organisationnelle relève ici de l'analyse de risque du chapitre 5.",
          colonnes_sev=(2,), largeurs=(0.7, 2, 1, 2.8))

    soa_donnees = (steps.get("evaluation") or {}).get("soa") or []
    if soa_donnees:
        _sous_titre(doc, "Déclaration d'Applicabilité (SoA) — synthèse par thème")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _run(p, "Détail des 93 contrôles de l'Annexe A dans le livrable dédié "
               "« Déclaration d'Applicabilité ».", couleur=_DOUX, taille=8.5, italique=True)
        _table(doc, ("Thème", "Total", "Applicables", "Exclus", "Non statués"),
              [(t["theme"], t["total"], t["applicables"], t["exclus"], t["non_statues"])
               for t in soa_module.par_theme(soa_donnees)],
              "", colonnes_num=(1, 2, 3, 4), largeurs=(1.6, 0.7, 0.9, 0.7, 0.9))


def _ch_technique(doc: Document, state: dict, steps: dict) -> None:
    resultat = couverture.couverture_technique(state)
    _encadre(doc, couverture.phrase(resultat))
    technique = (steps.get("evaluation") or {}).get("technical_results") or {}
    if not technique:
        _vide(doc, "Aucun scan technique de configuration n'a été exécuté pour cette mission.")
        return
    _table(doc, ("Indicateur", "Valeur"),
          [("Score technique", f"{technique.get('score')} % ({technique.get('band')})"),
           ("Failles critiques", technique.get("critical_count"))],
          "Résultats techniques indisponibles.", colonnes_num=(1,),
          largeurs=(2, 2))


def _ch_rattachement(doc: Document, state: dict, steps: dict) -> None:
    resultat = controles_techniques.etat(state)
    lignes = [(p["libelle"],
               ", ".join(f"{m['referentiel']} {m['ref']}" for m in p["mappings"]),
               "Couverte" if p["couverte"] else "Non couverte",
               f"{p['justification']} (phase {p['phase']})")
              for p in resultat["pratiques"]]
    _table(doc, ("Pratique", "Contrôles rattachés", "État", "Constaté en"), lignes,
          "Aucune pratique n'est rattachée.",
          largeurs=(1.8, 1.6, 1, 2.1))
    _note(doc, f"{resultat['couvertes']} pratique(s) couverte(s) sur "
              f"{resultat['total']} — {resultat['taux']} %.")


def _ch_traitement(doc: Document, state: dict, steps: dict, prefixe: str = "11") -> None:
    """Réutilisé tel quel par `build_ebios_docx` (prefixe="4").

    Régression corrigée le 31/07/2026 : « Plan de traitement » est le
    chapitre 11 (voir `report_html.CHAPITRES`), mais ces sous-titres
    affichaient « 10.1 »/« 10.2 » depuis l'ajout du chapitre AIPD, qui a
    décalé toute la numérotation en aval sans que ces deux chaînes codées en
    dur ne suivent — même bug que dans `report_html.py`, corrigé en miroir.
    """
    remediations = (steps.get("traitement") or {}).get("remediations") or []
    ordre = {"Critique": 0, "Élevé": 1, "Moyen": 2, "Faible": 3}
    triees = sorted(remediations, key=lambda r: ordre.get(r.get("priority"), 9))
    _sous_titre(doc, f"{prefixe}.1 Mesures priorisées")
    _table(doc, ("ID", "Priorité", "Axe", "Mesure de traitement"),
          [(r.get("id"), r.get("priority"), r.get("axe"), r.get("measure")) for r in triees],
          "Aucune mesure de traitement n'a été définie à ce stade.", colonnes_sev=(1,),
          largeurs=(0.7, 1, 1.1, 2.7))

    _sous_titre(doc, f"{prefixe}.1bis Pilotage (responsable, échéance, statut)")
    _table(doc, ("ID", "Responsable", "Échéance", "Statut", "Coût estimé"),
          [(r.get("id"), r.get("responsable"), r.get("echeance"), r.get("statut"), r.get("cout_estime"))
           for r in triees],
          "Aucune mesure de traitement n'a été définie à ce stade.",
          largeurs=(0.6, 1.4, 1, 1, 1.5))

    _sous_titre(doc, f"{prefixe}.2 Actions immédiates")
    wins = (steps.get("traitement") or {}).get("quick_wins") or []
    if not wins:
        _vide(doc, "Aucune action immédiate n'a été retenue.")
        return
    for i, w in enumerate(wins, 1):
        p = doc.add_paragraph(style="List Number")
        _run(p, _t(w), couleur=_CORPS, taille=9.5)


def _ch_charges(doc: Document, state: dict, steps: dict) -> None:
    socle = state.get("socle") or {}
    entrees = ((socle.get("temps") or {}).get("entrees")) or []
    budget = ((socle.get("qualification") or {}).get("budget")) or ""
    if not entrees:
        _vide(doc, "Aucun temps consommé n'a été saisi pour cette mission.")
        return

    from . import report_builder  # noqa: PLC0415 — libellés de durée partagés

    par_phase: dict[str, int] = {}
    for e in entrees:
        cle = e.get("phase", "autre")
        par_phase[cle] = par_phase.get(cle, 0) + int(e.get("minutes") or 0)
    lignes = [(libelle, report_builder._duree_lisible(par_phase[cle]))
              for cle, libelle in report_builder.PHASES_LIBELLES.items() if par_phase.get(cle)]
    lignes.append(("Total", report_builder._duree_lisible(sum(par_phase.values()))))
    _table(doc, ("Phase", "Temps consommé"), lignes, "Aucun temps consommé.", colonnes_num=(1,),
          largeurs=(2.2, 1.8))
    if budget:
        _note(doc, f"Budget vendu : {budget}")


def _ch_reserves(doc: Document, state: dict, steps: dict) -> None:
    date_emission = datetime.now().strftime("%d/%m/%Y")
    texte = docx_export.mention_reserve(date_emission, str(state.get("client") or ""))
    p = doc.add_paragraph()
    _run(p, texte, couleur=_CORPS, taille=9.5)


_BATISSEURS = (
    _ch_synthese, _ch_cadrage, _ch_patrimoine, _ch_aipd, _ch_risque, _ch_ecosysteme,
    _ch_resilience, _ch_evaluation, _ch_technique, _ch_rattachement, _ch_traitement,
    _ch_charges, _ch_reserves,
)


# --- Page de garde, sommaire, pied de page -----------------------------------

def _page_de_garde(doc: Document, titre: str, client: str, mission: str,
                   meta: list[tuple[str, str]], cabinet: str = "", logo: str = "") -> None:
    section = doc.sections[0]
    bande = doc.add_table(rows=1, cols=1)
    ligne = bande.rows[0]
    # Bandeau plein page plutôt que collé en haut : sans quoi son contenu
    # (logo + quelques lignes) laisse un grand vide blanc au-dessus du pied
    # de page sur toute page de garde courte — repéré le 30/07/2026 sur le
    # NDA (une seule page de garde très légère). `AT_LEAST` laisse le
    # bandeau grandir si le contenu (nom de mission long, etc.) l'exige.
    hauteur_utile = section.page_height - section.top_margin - section.bottom_margin
    ligne.height = Emu(int(hauteur_utile) - Cm(0.6))
    ligne.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cell = ligne.cells[0]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _shade(cell, _SOMBRE)
    cell.margin_top = Cm(1.2)
    cell.margin_bottom = Cm(1.2)
    cell.margin_left = Cm(1.2)
    cell.margin_right = Cm(1.2)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(_logo_bytes(logo), width=Cm(1.9))

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    _run(p, "GREEN SHIELD", gras=True, couleur="FFFFFF", taille=13)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p, cabinet or _CABINET_DEFAUT, italique=True, couleur=_VERT_CLAIR, taille=9)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    _run(p, titre, gras=True, couleur="FFFFFF", taille=22)

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    _run(p, client, gras=True, couleur=_VERT, taille=14)

    if mission:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, mission, couleur="A8C6BC", taille=10)

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    for i, (label, valeur) in enumerate(meta):
        if i:
            p.add_run("    ·    ").font.color.rgb = _rgb("3E5A50")
        _run(p, f"{label} : ", couleur="8FB3A8", taille=8.5)
        _run(p, str(valeur), gras=True, couleur="EAF4F0", taille=8.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    _run(p, "  DOCUMENT CONFIDENTIEL — DIFFUSION RESTREINTE  ", gras=True, couleur=_VERT_CLAIR, taille=8)

    _fixer_largeurs(bande, [int(_largeur_utile(doc))])
    doc.add_page_break()


def _sommaire(doc: Document, titres: list[str]) -> None:
    """Bloc sommaire générique : une entrée numérotée par titre de `titres`.

    Chaque document compose sa propre liste (le rapport de mission y ajoute
    « Certifications et signatures » ; les livrables plus courts n'en ont pas
    besoin — un NDA d'une page n'a pas de sommaire).
    """
    p = doc.add_paragraph()
    _run(p, "Sommaire", gras=True, couleur=_ENCRE, taille=13)
    p.paragraph_format.space_after = Pt(8)

    for numero, nom in enumerate(titres, 1):
        ligne = doc.add_paragraph()
        ligne.paragraph_format.space_after = Pt(2)
        _run(ligne, f"{numero}.  ", couleur=_VERT, taille=10)
        _run(ligne, nom, couleur=_CORPS, taille=10)

    doc.add_page_break()


def _signatures(doc: Document, numero: int, auditeur: str, cabinet: str, client: str) -> None:
    _chapitre_titre(doc, numero, "Certifications et signatures")
    p = doc.add_paragraph()
    _run(p, "L'auditeur certifie l'exactitude des constats factuels mentionnés dans le "
           "présent rapport.", couleur=_CORPS, taille=9.5)

    table = doc.add_table(rows=1, cols=2)
    entetes = table.rows[0].cells
    for i, texte in enumerate(("Signature de l'auditeur", "Signature du client audité")):
        _shade(entetes[i], _TEINTE)
        _run(entetes[i].paragraphs[0], texte, gras=True, couleur=_ENCRE, taille=8)

    contenu = table.add_row().cells
    p0 = contenu[0].paragraphs[0]
    _run(p0, auditeur or "—", gras=True, couleur=_CORPS, taille=9)
    p0b = contenu[0].add_paragraph()
    _run(p0b, cabinet or _CABINET_DEFAUT, couleur=_DOUX, taille=9)

    p1 = contenu[1].paragraphs[0]
    _run(p1, "DSI / Responsable de la sécurité", gras=True, couleur=_CORPS, taille=9)
    p1b = contenu[1].add_paragraph()
    _run(p1b, client, couleur=_DOUX, taille=9)

    _fixer_largeurs(table, _largeurs_ponderees(doc, (1, 1)))


def _pied_de_page(section, empreinte: str, cabinet: str = "") -> None:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Couleur du corps de texte (`_CORPS`, ~11:1 sur blanc), pas `_DOUX`
    # (~4,3:1) : à une taille aussi petite (6,5-7 pt), même un contraste
    # correct sur le papier finit visuellement pâle une fois combiné au
    # rendu grisé que Word applique par défaut à un pied de page inactif —
    # retour utilisateur du 30/07/2026 (« caractères presque blancs »). La
    # hiérarchie « ligne la plus discrète » reste portée par l'italique et
    # la taille, jamais par le contraste de couleur.
    _run(p, f"GREEN SHIELD — {cabinet or _CABINET_DEFAUT} · Document confidentiel, ne pas "
           "diffuser sans autorisation écrite.", couleur=_CORPS, taille=7)
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p2, f"Empreinte SHA-256 de l'état de la mission à l'édition : {empreinte}",
        couleur=_CORPS, taille=6.5)
    p3 = footer.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p3, "Toute modification ultérieure de la mission, même rétablie, produit une "
            "empreinte différente.", italique=True, couleur=_CORPS, taille=6.5)


def _nouveau_document() -> tuple[Document, object]:
    """Document Word vierge, format et style communs aux cinq livrables :
    A4 (pas le Letter par défaut de `python-docx` — le rapport HTML imprime en
    A4 et le contexte du projet est français/européen), corps en Calibri
    10,5 pt teinté `_CORPS`, marges identiques."""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = _rgb(_CORPS)

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.2)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.8)
    return doc, section


def build_report_docx(state: dict, p_id: str, auditeur: str = "", cabinet: str = "", logo: str = "") -> tuple[str, bytes]:
    """Rend le rapport de mission complet au format Word. Retourne (nom, octets)."""
    est_grc, titre, meta = report_html.titre_et_meta(state, p_id, auditeur, cabinet)
    client = str(state.get("client") or "")
    mission_nom = str(state.get("name") or "")
    steps = state.get("steps") or {}
    empreinte = docx_export.data_fingerprint(state)

    doc, section = _nouveau_document()
    _page_de_garde(doc, titre, client, mission_nom, meta, cabinet=cabinet, logo=logo)
    _sommaire(doc, [nom for nom, _rendu in report_html.CHAPITRES] + ["Certifications et signatures"])

    for numero, ((nom, _rendu), batisseur) in enumerate(zip(report_html.CHAPITRES, _BATISSEURS), 1):
        _chapitre_titre(doc, numero, nom)
        batisseur(doc, state, steps)

    _signatures(doc, len(report_html.CHAPITRES) + 1, auditeur, cabinet, client)
    _pied_de_page(section, empreinte, cabinet=cabinet)

    buffer = io.BytesIO()
    doc.save(buffer)
    nom_fichier = f"Rapport_{'GRC' if est_grc else 'Conseil'}_{p_id}.docx"
    return nom_fichier, buffer.getvalue()


# --- Les quatre autres livrables — même identité visuelle que le rapport ----
#
# Chacune de ces fonctions restitue exactement le même contenu que la version
# Markdown correspondante (`report_builder.py::build_document`, doc_type
# "nda"/"ebios"/"pssi_pri"/"aipd") : même source de données, même structure,
# seule la mise en forme change. Là où le contenu est identique à une section
# du rapport de mission (patrimoine, cartographie des menaces, écosystème des
# tiers, plan d'action), les bâtisseurs `_ch_*` sont réutilisés avec leur
# propre numérotation plutôt que dupliqués.

def build_nda_docx(state: dict, p_id: str, auditeur: str = "", cabinet: str = "", logo: str = "") -> tuple[str, bytes]:
    """Accord de confidentialité (NDA) — même contenu que `report_builder.py`
    (doc_type="nda")."""
    client = str(state.get("client") or "")
    steps = state.get("steps") or {}
    empreinte = docx_export.data_fingerprint(state)
    now = datetime.now().strftime("%d/%m/%Y")
    nda_text = (steps.get("cadrage") or {}).get("nda_text") or "NDA non rédigé."

    meta = [("Réf. mission", p_id), ("Édité le", now),
            ("Auditeur", auditeur or "—"), ("Cabinet", cabinet or _CABINET_DEFAUT)]

    doc, section = _nouveau_document()
    _page_de_garde(doc, "Accord de confidentialité", client, str(state.get("name") or ""), meta, cabinet=cabinet, logo=logo)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    _run(p, "CLASSIFICATION : CONFIDENTIEL AFFAIRES", gras=True, couleur=_ROSE, taille=9)

    for paragraphe in nda_text.split("\n"):
        if paragraphe.strip():
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(8)
            _run(para, paragraphe.strip(), couleur=_CORPS, taille=10.5)

    _titre_sans_numero(doc, "Signatures")
    p = doc.add_paragraph()
    _run(p, "En foi de quoi, les parties s'engagent et signent électroniquement ou de manière "
           "manuscrite :", couleur=_CORPS, taille=9.5)
    p.paragraph_format.space_after = Pt(6)
    _table(doc, (f"Pour {cabinet or _CABINET_DEFAUT}", f"Pour {client}"),
          [(f"{auditeur or _AUDITEUR_DEFAUT}, Consultant Cyber", "Mandataire habilité"),
           (f"Signature cryptographique locale : SHA256:{empreinte}", "Signature :"),
           (f"Date : {now}", "Date :")],
          "—", largeurs=(1, 1))

    _pied_de_page(section, empreinte, cabinet=cabinet)
    buffer = io.BytesIO()
    doc.save(buffer)
    return f"Accord_Confidentialite_{p_id}.docx", buffer.getvalue()


def build_ebios_docx(state: dict, p_id: str, auditeur: str = "", cabinet: str = "", logo: str = "") -> tuple[str, bytes]:
    """Analyse de risques EBIOS RM — même contenu que `report_builder.py`
    (doc_type="ebios"). Réutilise les bâtisseurs de patrimoine, de cartographie
    des menaces, d'écosystème et de plan d'action du rapport de mission, avec
    leur propre numérotation (1 à 4), indépendante de celle du rapport."""
    client = str(state.get("client") or "")
    steps = state.get("steps") or {}
    empreinte = docx_export.data_fingerprint(state)
    now = datetime.now().strftime("%d/%m/%Y")

    meta = [("Réf. mission", p_id), ("Méthode", "EBIOS RM · ANSSI"),
            ("Auditeur", auditeur or "—"), ("Cabinet", cabinet or _CABINET_DEFAUT),
            ("Édité le", now)]
    titres = ("Cadrage et identification du patrimoine",
              "Cartographie des menaces & scénarios EBIOS RM",
              "Écosystème et risques tiers",
              "Plan d'action & traitement")

    doc, section = _nouveau_document()
    _page_de_garde(doc, "Analyse de risques cyber (orientation EBIOS RM)", client,
                   str(state.get("name") or ""), meta, cabinet=cabinet, logo=logo)
    _sommaire(doc, list(titres))

    _chapitre_titre(doc, 1, titres[0])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _run(p, "Ce chapitre identifie le périmètre d'évaluation, les missions fondamentales de "
           "l'entreprise et cartographie le patrimoine d'actifs.", couleur=_CORPS, taille=10)
    _ch_patrimoine(doc, state, steps, prefixe="1")

    _chapitre_titre(doc, 2, titres[1])
    _ch_risque(doc, state, steps, prefixe="2")

    _chapitre_titre(doc, 3, titres[2])
    _ch_ecosysteme(doc, state, steps)

    _chapitre_titre(doc, 4, titres[3])
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _run(p, "Chaque mesure ci-dessous répond à un scénario ou à un écart constaté au "
           "chapitre 2.", couleur=_CORPS, taille=10)
    _ch_traitement(doc, state, steps, prefixe="4")

    _pied_de_page(section, empreinte, cabinet=cabinet)
    buffer = io.BytesIO()
    doc.save(buffer)
    return f"Analyse_Risques_EBIOS_{p_id}.docx", buffer.getvalue()


def build_pssi_docx(state: dict, p_id: str, auditeur: str = "", cabinet: str = "", logo: str = "") -> tuple[str, bytes]:
    """PSSI & Plan de reprise (PRI) — même contenu que `report_builder.py`
    (doc_type="pssi_pri")."""
    client = str(state.get("client") or "")
    steps = state.get("steps") or {}
    empreinte = docx_export.data_fingerprint(state)
    now = datetime.now().strftime("%d/%m/%Y")
    resilience = steps.get("resilience") or {}
    bcp = resilience.get("bcp_strategy") or {}
    e3r = resilience.get("e3r") or {}
    strategie = resilience.get("strategie_remediation") or {}
    sections_pssi = (steps.get("pssi_pri") or {}).get("pssi_sections") or []

    meta = [("Réf. mission", p_id), ("Auditeur", auditeur or "—"),
            ("Cabinet", cabinet or _CABINET_DEFAUT), ("Édité le", now)]
    titres = ("I. Politique de sécurité de l'information (PSSI)",
              "II. Plan de reprise informatique & résilience (PRI)")

    doc, section = _nouveau_document()
    _page_de_garde(doc, "Politique de sécurité de l'information (PSSI) & Plan de reprise (PRI)",
                   client, str(state.get("name") or ""), meta, cabinet=cabinet, logo=logo)
    _sommaire(doc, list(titres))

    _chapitre_titre(doc, 1, titres[0])
    if not sections_pssi:
        # La check-list PSSI (Phase 2) n'a pas encore été rédigée : le dire
        # plutôt que de laisser un chapitre I muet, sans titre ni contenu.
        _vide(doc, "Aucune section PSSI n'a été rédigée à ce stade.")
    for s in sections_pssi:
        _sous_titre(doc, str(s.get("title") or "—"))
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _run(p, str(s.get("content") or ""), couleur=_CORPS, taille=10)

    _chapitre_titre(doc, 2, titres[1])
    _sous_titre(doc, "2.1 Indicateurs temporels de continuité")
    _table(doc, ("Indicateur", "Valeur retenue"),
          [(lib, val) for lib, val in (
              ("RTO — temps de reprise maximal", bcp.get("rto")),
              ("RPO — perte de données maximale admissible", bcp.get("rpo")),
          ) if str(val or "").strip()],
          "Aucun indicateur temporel n'a été défini.", largeurs=(2.6, 1.4))

    _sous_titre(doc, "2.2 Politique de sauvegarde et d'immuabilité")
    if str(bcp.get("backup_policy") or "").strip():
        p = doc.add_paragraph()
        _run(p, bcp["backup_policy"], couleur=_CORPS, taille=10)
    else:
        _vide(doc, "Aucune politique de sauvegarde n'a été définie.")

    _sous_titre(doc, "2.3 Séquence de remédiation en gestion de crise (E3R de l'ANSSI)")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _run(p, "En cas de compromission majeure de l'Active Directory ou de l'infrastructure "
           "cloud :", couleur=_DOUX, taille=9)
    _table(doc, ("Étape E3R", "Procédure retenue"),
          [(lib, e3r.get(cle)) for lib, cle in (
              ("Endiguement (contenir l'attaquant)", "endiguement"),
              ("Éviction (reprendre le contrôle du cœur de confiance)", "eviction"),
              ("Éradication (nettoyage en profondeur)", "eradication"),
              ("Reconstruction (rebâtir dès la conception)", "reconstruction"),
          ) if str(e3r.get(cle) or "").strip()],
          "La séquence de remédiation E3R n'a pas été documentée.", largeurs=(1.6, 3.4))

    _sous_titre(doc, "2.4 Volet stratégique — arbitrage Direction")
    _table(doc, ("Critère d'arbitrage", "Position retenue"),
          [(lib, strategie.get(cle)) for lib, cle in (
              ("Urgence de redémarrage", "urgence_redemarrage"),
              ("Coûts et risques d'un redémarrage précipité", "couts_risques_redemarrage"),
              ("Décision retenue et autorité", "decision_direction"),
          ) if str(strategie.get(cle) or "").strip()],
          "Le volet stratégique (arbitrage Direction) n'a pas été documenté.", largeurs=(1.8, 3.2))

    _titre_sans_numero(doc, "Signatures pour homologation de sécurité")
    _table(doc, (f"Pour {cabinet or _CABINET_DEFAUT}", f"Pour la Direction de {client}"),
          [(auditeur or _AUDITEUR_DEFAUT, "Directeur Général / RSSI"), ("Signature :", "Signature :")],
          "—", largeurs=(1, 1))

    _pied_de_page(section, empreinte, cabinet=cabinet)
    buffer = io.BytesIO()
    doc.save(buffer)
    return f"PSSI_PRI_{p_id}.docx", buffer.getvalue()


def build_aipd_docx(state: dict, p_id: str, auditeur: str = "", cabinet: str = "", logo: str = "") -> tuple[str, bytes]:
    """AIPD / PIA (RGPD) — même contenu que `report_builder.py` (doc_type="aipd").

    Ne conditionne pas son contenu à `aipd_required`, à l'identique de la
    version Markdown : ce document n'est produit que sur demande explicite du
    consultant (bouton dédié), qui a pu juger utile une AIPD volontaire même
    hors des cas où l'application la signale requise.
    """
    client = str(state.get("client") or "")
    steps = state.get("steps") or {}
    empreinte = docx_export.data_fingerprint(state)
    now = datetime.now().strftime("%d/%m/%Y")
    diagnostic = steps.get("diagnostic") or {}
    donnees = diagnostic.get("aipd") or {}

    meta = [("Réf. mission", p_id), ("DPO", "Enregistré au registre"),
            ("Auditeur", auditeur or "—"), ("Cabinet", cabinet or _CABINET_DEFAUT),
            ("Édité le", now)]
    titres = ("Registre des activités de traitement (RGPD Art. 30)",
              "Analyse d'impact systématique (PIA)",
              "Obligations organisationnelles (conduite de l'AIPD)")

    doc, section = _nouveau_document()
    _page_de_garde(doc, "Analyse d'impact relative à la protection des données (AIPD / PIA)",
                   client, str(state.get("name") or ""), meta, cabinet=cabinet, logo=logo)
    _sommaire(doc, list(titres))

    _chapitre_titre(doc, 1, titres[0])
    registre = diagnostic.get("rgpd_register") or []
    _table(doc, ("ID", "Activité de traitement", "Finalité", "Catégories de données", "Conservation"),
          [(r.get("id"), r.get("name"), r.get("purpose"), r.get("data_categories"),
            r.get("retention")) for r in registre],
          "Aucun traitement n'a été inscrit au registre.", largeurs=(0.7, 1.6, 2, 1.8, 1.3))

    _sous_titre(doc, "1.bis Registre des violations de données (Art. 33-34)")
    _table(doc, ("ID", "Constatée le", "Nature", "CNIL", "Personnes informées"),
          [(v.get("id"), v.get("date_constat"), v.get("nature"),
            f"Notifiée le {v.get('date_notification_cnil')}" if v.get("notifiee_cnil") else "Non notifiée",
            "Oui" if v.get("personnes_informees") else "Non")
           for v in diagnostic.get("violations") or []],
          "Aucune violation de données n'a été constatée sur cette mission.",
          largeurs=(0.7, 1, 1.6, 1.4, 1.1))

    _chapitre_titre(doc, 2, titres[1])
    for numero_volet, (sous, cle) in enumerate((
        ("Description systématique du traitement", "treatment_description"),
        ("Évaluation de la nécessité et de la proportionnalité", "necessity_eval"),
        ("Évaluation des risques sur les droits et libertés des personnes", "risks_eval"),
        ("Mesures de traitement et de sécurité envisagées (atténuation)", "mitigation_measures"),
    ), 1):
        _sous_titre(doc, f"2.{numero_volet} {sous}")
        contenu = str(donnees.get(cle) or "").strip()
        if contenu:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            _run(p, contenu, couleur=_CORPS, taille=10)
        else:
            _vide(doc, f"Le volet « {sous} » n'a pas été rédigé.")

    _chapitre_titre(doc, 3, titres[2])
    _table_obligations_aipd(doc, donnees)

    _titre_sans_numero(doc, "Signature de validation conformité CNIL")
    _table(doc, ("Avis du Délégué à la Protection des Données (DPO)",
                "Validation du Responsable du Traitement"),
          [("Avis favorable / non favorable", "Validé pour mise en œuvre"),
           ("Signature :", "Signature :")],
          "—", largeurs=(1, 1))

    _pied_de_page(section, empreinte, cabinet=cabinet)
    buffer = io.BytesIO()
    doc.save(buffer)
    return f"AIPD_RGPD_{p_id}.docx", buffer.getvalue()


def _libelle_applicable(valeur) -> str:
    if valeur is True:
        return "Applicable"
    if valeur is False:
        return "Exclu"
    return "Non statué"


def build_soa_docx(state: dict, p_id: str, auditeur: str = "", cabinet: str = "", logo: str = "") -> tuple[str, bytes]:
    """Déclaration d'Applicabilité (SoA) — ISO/IEC 27001:2022 Annexe A.

    Manque identifié en revue GRC senior le 30/07/2026 : sans SoA, une
    mission ISO 27001 ne peut pas passer un audit de certification, c'est le
    premier document qu'un auditeur externe demande (clause 6.1.3 d).

    Ne se produit que si `steps.evaluation.soa` existe (mission au
    référentiel ISO 27001) — la route appelante le vérifie et répond 404
    sinon, plutôt que de générer un document vide et trompeur.
    """
    client = str(state.get("client") or "")
    steps = state.get("steps") or {}
    empreinte = docx_export.data_fingerprint(state)
    now = datetime.now().strftime("%d/%m/%Y")
    donnees = (steps.get("evaluation") or {}).get("soa") or []
    resume = soa_module.etat(donnees)

    meta = [("Réf. mission", p_id), ("Référentiel", "ISO/IEC 27001:2022 Annexe A"),
            ("Auditeur", auditeur or "—"), ("Cabinet", cabinet or _CABINET_DEFAUT),
            ("Édité le", now)]

    doc, section = _nouveau_document()
    _page_de_garde(doc, "Déclaration d'Applicabilité (SoA)", client,
                   str(state.get("name") or ""), meta, cabinet=cabinet, logo=logo)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    _run(p, f"{resume['statues']}/{resume['total']} contrôle(s) statué(s) ({resume['taux']} %) — "
           f"{resume['applicables']} applicable(s), {resume['exclus']} exclu(s).",
        couleur=_CORPS, taille=9.5, gras=True)

    for numero, theme in enumerate(soa_module.THEMES, 1):
        entrees = [e for e in donnees if e.get("theme") == theme]
        _chapitre_titre(doc, numero, theme)
        _table(doc, ("Code", "Contrôle", "Applicabilité", "Statut", "Justification"),
              [(e.get("code"), e.get("titre"), _libelle_applicable(e.get("applicable")),
                e.get("statut") or "—", e.get("justification")) for e in entrees],
              f"Aucun contrôle du thème {theme}.",
              largeurs=(0.7, 1.9, 0.9, 0.9, 1.6))

    _pied_de_page(section, empreinte, cabinet=cabinet)
    buffer = io.BytesIO()
    doc.save(buffer)
    return f"Declaration_Applicabilite_SoA_{p_id}.docx", buffer.getvalue()
