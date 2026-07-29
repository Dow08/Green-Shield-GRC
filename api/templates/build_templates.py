"""Génère les gabarits Word (.docx) porteurs de balises Jinja2 pour docxtpl.

Pourquoi un script plutôt qu'un binaire committé tel quel : le gabarit doit être
reproductible et son contenu doit être lisible/versionnable en clair. Une fois
généré, le .docx est un document Word ordinaire — Dorian peut l'ouvrir, changer
les polices, les couleurs, l'en-tête, ajouter le logo, sans toucher au code
Python. Seules les balises {{ ... }} et {%tr ... %} doivent être préservées.

Usage :
    py -3 api/templates/build_templates.py            # ne réécrit pas un gabarit existant
    py -3 api/templates/build_templates.py --force    # réécrit (ÉCRASE les retouches Word)
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

TEMPLATES_DIR = Path(__file__).resolve().parent
ISO27001_TEMPLATE = TEMPLATES_DIR / "rapport_iso27001.docx"
# Logo servi à l'application web : source unique de l'identité visuelle, pour
# que le rapport Word et les livrables Markdown portent exactement la même.
LOGO_PATH = TEMPLATES_DIR.parent.parent / "web" / "public" / "logo.png"


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Petit tableau clé/valeur (informations de mission, traçabilité)."""
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = value


def _loop_table(doc: Document, headers: list[str], loop_var: str,
                collection: str, fields: list[str]) -> None:
    """Tableau dont les lignes sont répétées par docxtpl.

    Une balise {%tr ... %} fait disparaître la LIGNE qui la contient, remplacée
    par la directive Jinja correspondante. Les deux balises ne peuvent donc pas
    cohabiter dans la même ligne : il faut trois lignes distinctes — l'ouverture
    de boucle, la ligne modèle réellement répétée, puis la fermeture.
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, head in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = head
        cell.paragraphs[0].runs[0].bold = True

    table.add_row().cells[0].text = f"{{%tr for {loop_var} in {collection} %}}"

    content = table.add_row().cells
    for i, field in enumerate(fields):
        content[i].text = f"{{{{ {loop_var}.{field} }}}}"

    table.add_row().cells[0].text = "{%tr endfor %}"


def build_iso27001_template(path: Path) -> None:
    doc = Document()

    # Corps de texte un peu plus dense que le défaut Word.
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # --- Page de garde ---
    if LOGO_PATH.is_file():
        logo_par = doc.add_paragraph()
        logo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_par.add_run().add_picture(str(LOGO_PATH), width=Inches(1.1))
        marque = doc.add_paragraph()
        marque.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_marque = marque.add_run("GREEN SHIELD")
        run_marque.bold = True
        run_marque.font.size = Pt(16)
        cabinet = doc.add_paragraph()
        cabinet.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_cab = cabinet.add_run("DP Cyber Consulting — Audit & Conseil Cybersécurité")
        run_cab.italic = True
        run_cab.font.size = Pt(9)

    title = doc.add_heading("{{ titre_rapport }}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("{{ client }} — {{ referentiel }}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _kv_table(doc, [
        ("Client", "{{ client }}"),
        ("Mission", "{{ mission }}"),
        ("Référentiel", "{{ referentiel }}"),
        ("Auditeur", "{{ auditeur }} — {{ cabinet }}"),
        ("Périmètre", "{{ perimetre }}"),
        ("Date d'émission", "{{ date_emission }}"),
        ("Version du document", "{{ version }}"),
    ])

    p = doc.add_paragraph()
    run = p.add_run(
        "Document confidentiel. Diffusion restreinte aux destinataires désignés "
        "par {{ client }}. Toute reproduction ou communication à un tiers requiert "
        "l'accord écrit préalable des parties."
    )
    run.italic = True

    doc.add_page_break()

    # --- 1. Synthèse exécutive ---
    doc.add_heading("1. Synthèse exécutive", level=1)
    doc.add_paragraph("{{ synthese_executive }}")
    _kv_table(doc, [
        ("Score de conformité", "{{ score }} %"),
        ("Niveau de maîtrise", "{{ band }}"),
        ("Écarts identifiés", "{{ nb_ecarts }}"),
        ("Dont critiques", "{{ nb_critiques }}"),
    ])

    # --- 2. Périmètre et méthodologie ---
    doc.add_heading("2. Périmètre et méthodologie", level=1)
    doc.add_paragraph("{{ methodologie }}")

    # --- 3. Patrimoine informationnel ---
    doc.add_heading("3. Patrimoine informationnel", level=1)
    doc.add_heading("3.1 Valeurs métier", level=2)
    _loop_table(
        doc,
        ["Réf.", "Valeur métier", "Description", "Données personnelles"],
        "vm", "valeurs_metier",
        ["id", "name", "description", "personal"],
    )
    doc.add_heading("3.2 Biens supports", level=2)
    _loop_table(
        doc,
        ["Réf.", "Bien support", "Type", "Propriétaire"],
        "bs", "biens_supports",
        ["id", "name", "type", "owner"],
    )

    # --- 4. Constats ---
    doc.add_heading("4. Constats d'audit", level=1)
    doc.add_paragraph(
        "Chaque constat ci-dessous est appuyé par une preuve collectée ou une "
        "déclaration attribuée, conformément à la norme ISO 19011."
    )
    _loop_table(
        doc,
        ["Réf.", "Constat", "Statut", "Sévérité", "Preuve"],
        "c", "constats",
        ["id", "title", "status", "severity", "evidence"],
    )

    # --- 5. Plan d'action ---
    doc.add_heading("5. Plan d'action", level=1)
    _loop_table(
        doc,
        ["Réf.", "Axe", "Mesure corrective", "Priorité"],
        "a", "actions",
        ["id", "axe", "measure", "priority"],
    )

    # --- 6. Réserves et limites ---
    # Mention de réserve : délimite la responsabilité du consultant (spec §13.4).
    doc.add_heading("6. Réserves et limites", level=1)
    doc.add_paragraph("{{ mention_reserve }}")

    # --- 7. Traçabilité ---
    doc.add_heading("7. Traçabilité du document", level=1)
    _kv_table(doc, [
        ("Empreinte SHA-256 des données source", "{{ hash_donnees }}"),
        ("Généré le", "{{ date_emission }}"),
        ("Généré par", "GREEN SHIELD — {{ cabinet }}"),
    ])

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--force", action="store_true",
                        help="réécrit un gabarit existant (écrase les retouches faites dans Word)")
    args = parser.parse_args()

    if ISO27001_TEMPLATE.exists() and not args.force:
        print(f"[SKIP] {ISO27001_TEMPLATE.name} existe déjà. "
              f"Utiliser --force pour l'écraser (les retouches Word seront perdues).")
        return 0

    build_iso27001_template(ISO27001_TEMPLATE)
    print(f"[OK] Gabarit généré : {ISO27001_TEMPLATE}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
