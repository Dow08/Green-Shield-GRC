"""Tests du rapport de mission au format Word (report_docx.py).

Remplace le 31/07/2026 un gabarit `docxtpl` statique resté figé à 7 sections
génériques (aucun TPRM, AIPD, E3R ni DORA) pendant que `report_html.py`
montait à 13, avec un titre écrit en dur, faux sur une mission de conseil.
Ce module reconstruit le document en `python-docx` directement ; ces tests
vérifient qu'il rend un document Word valide, complet, et strictement aligné
sur `report_html.py` — même titre, mêmes chapitres, dans le même ordre.
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from modules import report_docx, report_html  # noqa: E402


def _mission(**overrides) -> dict:
    base = {"id": "acme", "name": "Audit ISO 27001", "client": "ACME", "type": "consulting", "steps": {}}
    base.update(overrides)
    return base


def _mission_complete(volet: str = "consulting") -> dict:
    """Une mission riche : au moins une donnée par chapitre du rapport."""
    return {
        "id": "acme", "name": "Audit de sécurité", "client": "ACME", "type": volet,
        "progress": 100,
        "socle": {
            "qualification": {"declencheur": "Exigence du donneur d'ordre", "budget": "12 jours"},
            "contractualisation": {"perimetre_inclus": "SI de production",
                                   "livrables": ["Rapport d'audit"]},
            "kickoff": {"date": "2026-06-03", "participants": ["RSSI"]},
            "entretiens": [{"id": "ENT-01", "role": "RSSI", "date": "2026-06-05",
                            "synthese": "Confirme l'absence de segmentation OT/IT."}],
            "temps": {"entrees": [{"phase": "cadrage", "minutes": 120}]},
        },
        "steps": {
            "cadrage": {
                "assets_metier": [{"id": "VM-01", "name": "Fichier clients",
                                   "description": "Coordonnées et contrats", "is_personal_data": True}],
                "assets_support": [{"id": "BS-01", "name": "Active Directory",
                                    "type": "Logiciel", "owner": "Équipe Système"}],
                "framework_name": "ISO/IEC 27001:2022" if volet == "grc" else None,
            },
            "diagnostic": {
                "rgpd_register": [{"id": "RGPD-01", "name": "Paie", "purpose": "Gestion RH",
                                   "data_categories": "Identité", "retention": "5 ans"}],
                "aipd_required": True,
                "aipd": {"treatment_description": "Vidéoprotection", "necessity_eval": "n",
                         "risks_eval": "r", "mitigation_measures": "m",
                         "risque_residuel": "eleve",
                         "obligations": [{"id": "DPO", "satisfait": True, "commentaire": "Avis rendu"}]},
            },
            "ebios": {
                "redoute_events": [{"id": "ER-01", "event": "Chiffrement par rançongiciel",
                                    "gravity": 4, "impact": "Arrêt de production"}],
                "risk_sources": [{"id": "SR-01", "name": "Cybercriminels", "objective": "Extorsion"}],
                "operational_scenarios": [{"id": "SO-01", "event": "Hameçonnage puis rebond",
                                           "gravity": 4, "likelihood": 3, "mitigation": "MFA"}],
                "case_studies": [{"case": "Norsk Hydro", "lessons": "Segmenter OT/IT"}],
            },
            "tprm": {"tiers": [
                {"name": "Prestataire critique", "dependence": 5, "penetration": 5,
                 "maturity": 2, "trust": 3, "score": 4.17, "rating": "Critique",
                 "exigences": [{"id": "NIST-ID.RA-10", "libelle": "Évaluation avant acquisition",
                               "satisfait": False}]},
            ]},
            "resilience": {
                "bcp_strategy": {"rto": "4 h", "rpo": "1 h", "backup_policy": "Immuable"},
                "e3r": {"endiguement": "Isolement", "eviction": "Révocation",
                       "eradication": "Reconstruction", "reconstruction": "Restauration"},
                "strategie_remediation": {"decision_direction": "Priorité à l'éradication avant redémarrage"},
            },
            "evaluation": {"manual_controls": [{"id": "A.5.1", "title": "Politiques de sécurité",
                                                "status": "NON_CONFORME", "notes": "PSSI absente"}]},
            "traitement": {
                "remediations": [{"id": "REM-01", "axe": "Protection", "measure": "Déployer le MFA",
                                  "priority": "Critique"}],
                "quick_wins": ["Activer le MFA sur la messagerie"],
            },
            "restitution": {"exec_summary": "Deux écarts majeurs, aucun structurel."},
        },
    }


def _all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _footer_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.sections[0].footer.paragraphs)


# --- Le fichier produit -------------------------------------------------------

def test_rendu_est_un_fichier_word_valide():
    nom, contenu = report_docx.build_report_docx(_mission(), "acme")
    assert contenu[:2] == b"PK", "un .docx est une archive ZIP : signature PK attendue"
    Document(io.BytesIO(contenu))  # lève si l'archive n'est pas un document Word
    assert nom.endswith(".docx")


def test_le_nom_de_fichier_varie_selon_le_volet():
    nom_conseil, _ = report_docx.build_report_docx(_mission(type="consulting"), "acme")
    nom_grc, _ = report_docx.build_report_docx(_mission(type="grc"), "acme")
    assert "Conseil" in nom_conseil
    assert "GRC" in nom_grc


# --- Le titre suit le volet (c'est le bug d'origine) -------------------------

def test_le_titre_suit_le_volet_conseil():
    """C'est le bug constaté le 31/07/2026 : le gabarit affichait « Rapport
    d'audit de conformité » même sur une mission de conseil."""
    _, contenu = report_docx.build_report_docx(_mission(type="consulting"), "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "analyse de risque" in texte
    assert "Rapport d'audit de conformité & GRC" not in texte


def test_le_titre_suit_le_volet_grc():
    _, contenu = report_docx.build_report_docx(_mission(type="grc"), "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "conformité" in texte


def test_le_titre_est_identique_a_celui_du_rapport_html():
    """Les deux formats calculent leur titre via la même fonction
    (`report_html.titre_et_meta`) : ils ne peuvent plus diverger."""
    for volet in ("consulting", "grc"):
        mission = _mission(type=volet)
        _, titre_html, _ = report_html.titre_et_meta(mission, "acme", "Dorian", "DP Cyber")
        _, contenu = report_docx.build_report_docx(mission, "acme", "Dorian", "DP Cyber")
        assert titre_html in _all_text(Document(io.BytesIO(contenu)))


# --- Sommaire et chapitres alignés sur report_html.py ------------------------

def test_le_sommaire_reprend_tous_les_titres_de_report_html():
    """Garde-fou contre la dérive qui a produit le bug d'origine : si un
    chapitre est ajouté à report_html.py sans toucher à ce module, ce test
    échoue plutôt que de laisser le Word prendre du retard en silence."""
    _, contenu = report_docx.build_report_docx(_mission(), "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    for nom, _rendu in report_html.CHAPITRES:
        assert nom in texte, f"chapitre absent du Word : {nom!r}"
    assert "Certifications et signatures" in texte


def test_chaque_chapitre_a_son_batisseur():
    """Un chapitre ajouté à CHAPITRES sans bâtisseur correspondant planterait
    au rendu plutôt que de produire un chapitre vide en silence."""
    assert len(report_docx._BATISSEURS) == len(report_html.CHAPITRES)


# --- Contenu réellement restitué, chapitre par chapitre ----------------------

def test_toutes_les_donnees_de_la_mission_apparaissent():
    _, contenu = report_docx.build_report_docx(_mission_complete(), "acme", "Dorian", "DP Cyber")
    texte = _all_text(Document(io.BytesIO(contenu))) + "\n" + _footer_text(Document(io.BytesIO(contenu)))
    attendus = [
        "Deux écarts majeurs, aucun structurel",          # synthèse
        "Exigence du donneur d'ordre",                     # cadrage
        "RSSI",                                            # entretiens
        "Fichier clients",                                  # patrimoine
        "Active Directory",
        "Vidéoprotection",                                  # AIPD
        "Chiffrement par rançongiciel",                     # événements redoutés
        "Cybercriminels",                                   # sources de risque
        "Hameçonnage puis rebond",                          # scénarios
        "Norsk Hydro",                                      # cas réels
        "Prestataire critique",                             # écosystème tiers
        "Isolement",                                        # E3R
        "Priorité à l'éradication avant redémarrage",       # volet stratégique ANSSI
        "Politiques de sécurité",                           # évaluation organisationnelle
        "Déployer le MFA",                                  # plan de traitement
        "Activer le MFA sur la messagerie",                 # actions immédiates
        "Dorian",                                           # signatures
    ]
    for attendu in attendus:
        assert attendu in texte, f"donnée absente du rapport Word : {attendu!r}"


def test_un_champ_vide_affiche_un_tiret_pas_une_invention():
    mission = _mission_complete()
    mission["steps"]["cadrage"]["assets_metier"][0]["description"] = ""
    _, contenu = report_docx.build_report_docx(mission, "acme")
    doc = Document(io.BytesIO(contenu))
    table = next(t for t in doc.tables if any("Fichier clients" in c.text for r in t.rows for c in r.cells))
    ligne = next(r for r in table.rows if any("Fichier clients" in c.text for c in r.cells))
    assert any(c.text.strip() == "—" for c in ligne.cells)


def test_une_section_sans_donnee_affiche_un_message_explicite():
    """Une section sans ligne n'est jamais un tableau à en-têtes vide."""
    _, contenu = report_docx.build_report_docx(_mission(), "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Aucune valeur métier n'a été cartographiée." in texte
    assert "Aucun événement redouté n'a été caractérisé." in texte


def test_la_synthese_non_redigee_est_signalee():
    _, contenu = report_docx.build_report_docx(_mission(), "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "jamais produite automatiquement" in texte


# --- Volet GRC : §14.1bis — aucun score de risque ----------------------------

def test_le_volet_grc_n_affiche_aucun_score_de_risque():
    mission = _mission_complete("grc")
    _, contenu = report_docx.build_report_docx(mission, "acme")
    doc = Document(io.BytesIO(contenu))
    entetes = {c.text for t in doc.tables for c in t.rows[0].cells}
    assert "Ratio" not in entetes
    assert "aucun score de risque" in _all_text(doc)


def test_le_volet_conseil_classe_les_tiers_par_ratio():
    mission = _mission_complete("consulting")
    _, contenu = report_docx.build_report_docx(mission, "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "4.17" in texte


# --- Statuts lisibles, pas les valeurs internes ------------------------------

def test_les_statuts_bruts_ne_fuient_pas_dans_le_document():
    mission = _mission_complete()
    _, contenu = report_docx.build_report_docx(mission, "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Non conforme" in texte
    assert "NON_CONFORME" not in texte


# --- Traçabilité --------------------------------------------------------------

def test_le_pied_de_page_porte_l_empreinte():
    from modules import docx_export
    mission = _mission()
    empreinte = docx_export.data_fingerprint(mission)
    _, contenu = report_docx.build_report_docx(mission, "acme")
    doc = Document(io.BytesIO(contenu))
    assert empreinte in _footer_text(doc)


def test_le_pied_de_page_rappelle_la_confidentialite():
    _, contenu = report_docx.build_report_docx(_mission(), "acme")
    doc = Document(io.BytesIO(contenu))
    assert "confidentiel" in _footer_text(doc).lower()


def test_la_reserve_cite_le_client():
    _, contenu = report_docx.build_report_docx(_mission(client="Vernier Composites"), "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Vernier Composites" in texte
    assert "ne saurait" in texte and "garantie" in texte


# --- Robustesse aux caractères spéciaux --------------------------------------

def test_les_caracteres_speciaux_survivent_a_l_export():
    """python-docx écrit chaque valeur dans un run de texte natif — il n'y a
    plus de gabarit à substituer, donc plus de risque qu'un « & » saisi soit
    interprété comme le début d'une entité XML (régression de recette du
    29/07/2026 sur l'ancien gabarit `docxtpl`)."""
    mission = _mission_complete()
    mission["steps"]["cadrage"]["assets_metier"][0]["name"] = "Formulations R&D"
    mission["client"] = "ACME & Fils <SAS>"
    _, contenu = report_docx.build_report_docx(mission, "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Formulations R&D" in texte
    assert "ACME & Fils <SAS>" in texte


# --- Page de garde -------------------------------------------------------------

def test_la_page_de_garde_porte_le_logo():
    _, contenu = report_docx.build_report_docx(_mission(), "acme")
    doc = Document(io.BytesIO(contenu))
    assert len(doc.inline_shapes) >= 1


def _image_embarquee(contenu: bytes) -> bytes:
    doc = Document(io.BytesIO(contenu))
    image = next(p for p in doc.part.related_parts.values()
                 if p.content_type.startswith("image/"))
    return image.blob


def test_le_logo_personnalise_remplace_le_logo_green_shield():
    """Ajouté le 30/07/2026 : sans logo de cabinet déposé dans Réglages, la
    page de garde reste celle de GREEN SHIELD par défaut ; un logo valide
    fourni prend sa place."""
    import base64
    from modules import charte

    _, sans_logo = report_docx.build_report_docx(_mission(), "acme")
    assert _image_embarquee(sans_logo) == base64.b64decode(charte.LOGO_BASE64)

    logo_perso = ("iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY"
                  "42YAAAAASUVORK5CYII=")
    _, avec_logo = report_docx.build_report_docx(_mission(), "acme", logo=logo_perso)
    assert _image_embarquee(avec_logo) == base64.b64decode(logo_perso)


def test_la_page_de_garde_porte_le_client_et_la_mission():
    mission = _mission(client="Vernier Composites SAS", name="Audit EBIOS RM")
    _, contenu = report_docx.build_report_docx(mission, "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Vernier Composites SAS" in texte
    assert "Audit EBIOS RM" in texte


# --- Aucune ligne de tableau entièrement vide --------------------------------

def test_aucune_ligne_de_tableau_n_est_entierement_vide():
    """Régression : `doc.add_table(rows=2, ...)` réserve déjà une ligne vide,
    et y appeler ensuite `table.add_row()` en ajoute une troisième — la
    ligne 2 reste blanche pour toujours. Trouvé dans le tableau de
    signatures en inspectant le .docx généré, pas par un test."""
    _, contenu = report_docx.build_report_docx(_mission_complete(), "acme", "Dorian", "DP Cyber")
    doc = Document(io.BytesIO(contenu))
    for t in doc.tables:
        for r in t.rows:
            assert any(c.text.strip() for c in r.cells), \
                "ligne de tableau entièrement vide (probable rows= + add_row() en trop)"


# =====================================================================
# Les quatre autres livrables (NDA, EBIOS, PSSI/PRI, AIPD) — ajoutés le
# 30/07/2026 pour donner à ces documents la même identité Word que le
# rapport de mission (page de garde, sommaire, tableaux, pied de page).
# =====================================================================

_BATISSEURS_AUTRES = {
    "nda": report_docx.build_nda_docx,
    "ebios": report_docx.build_ebios_docx,
    "pssi": report_docx.build_pssi_docx,
    "aipd": report_docx.build_aipd_docx,
}


@pytest.mark.parametrize("cle", sorted(_BATISSEURS_AUTRES))
def test_chaque_livrable_est_un_fichier_word_valide(cle):
    batisseur = _BATISSEURS_AUTRES[cle]
    nom, contenu = batisseur(_mission_complete(), "acme")
    assert contenu[:2] == b"PK"
    Document(io.BytesIO(contenu))
    assert nom.endswith(".docx")


@pytest.mark.parametrize("cle", sorted(_BATISSEURS_AUTRES))
def test_chaque_livrable_porte_l_empreinte_en_pied_de_page(cle):
    from modules import docx_export
    mission = _mission_complete()
    empreinte = docx_export.data_fingerprint(mission)
    _, contenu = _BATISSEURS_AUTRES[cle](mission, "acme")
    doc = Document(io.BytesIO(contenu))
    assert empreinte in _footer_text(doc)


@pytest.mark.parametrize("cle", sorted(_BATISSEURS_AUTRES))
def test_chaque_livrable_cite_le_client_sur_la_page_de_garde(cle):
    mission = _mission_complete()
    mission["client"] = "Vernier Composites SAS"
    _, contenu = _BATISSEURS_AUTRES[cle](mission, "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Vernier Composites SAS" in texte


def test_le_nom_de_fichier_de_chaque_livrable():
    assert report_docx.build_nda_docx(_mission(), "acme")[0] == "Accord_Confidentialite_acme.docx"
    assert report_docx.build_ebios_docx(_mission(), "acme")[0] == "Analyse_Risques_EBIOS_acme.docx"
    assert report_docx.build_pssi_docx(_mission(), "acme")[0] == "PSSI_PRI_acme.docx"
    assert report_docx.build_aipd_docx(_mission(), "acme")[0] == "AIPD_RGPD_acme.docx"


# --- NDA ----------------------------------------------------------------------

def test_nda_reprend_le_texte_redige_et_signale_son_absence():
    mission = _mission()
    mission["steps"]["cadrage"] = {"nda_text": "Entre les parties Dorian et ACME, il est convenu..."}
    _, contenu = report_docx.build_nda_docx(mission, "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Entre les parties Dorian et ACME" in texte

    _, contenu_vide = report_docx.build_nda_docx(_mission(), "acme")
    assert "NDA non rédigé." in _all_text(Document(io.BytesIO(contenu_vide)))


def test_nda_porte_un_tableau_de_signatures_avec_l_empreinte():
    from modules import docx_export
    mission = _mission()
    empreinte = docx_export.data_fingerprint(mission)
    _, contenu = report_docx.build_nda_docx(mission, "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Signatures" in texte
    assert empreinte in texte  # signature cryptographique locale, dans le corps du tableau


# --- EBIOS RM -------------------------------------------------------------------

def test_ebios_reutilise_les_donnees_de_patrimoine_et_de_menaces():
    """Ce document ne duplique pas les données : il rend les mêmes bâtisseurs
    que le rapport de mission (patrimoine, menaces, écosystème, traitement),
    avec sa propre numérotation 1 à 4."""
    mission = _mission_complete()
    _, contenu = report_docx.build_ebios_docx(mission, "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    for attendu in ("Fichier clients", "Active Directory", "Chiffrement par rançongiciel",
                    "Cybercriminels", "Hameçonnage puis rebond", "Norsk Hydro",
                    "Prestataire critique", "Déployer le MFA"):
        assert attendu in texte, f"donnée absente de l'EBIOS Word : {attendu!r}"


def test_ebios_numerote_ses_chapitres_de_un_a_quatre_independamment_du_rapport():
    _, contenu = report_docx.build_ebios_docx(_mission_complete(), "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "1. Cadrage et identification du patrimoine" in texte
    assert "2. Cartographie des menaces & scénarios EBIOS RM" in texte
    assert "3. Écosystème et risques tiers" in texte
    assert "4. Plan d'action & traitement" in texte


def test_ebios_n_a_pas_de_bloc_signatures():
    """À la différence du NDA et du PSSI, l'analyse EBIOS RM n'engage aucune
    signature — c'est un document d'analyse, pas un contrat."""
    _, contenu = report_docx.build_ebios_docx(_mission_complete(), "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Signatures" not in texte


# --- PSSI & PRI -----------------------------------------------------------------

def test_pssi_rend_les_sections_redigees_en_phase_2():
    mission = _mission()
    mission["steps"]["pssi_pri"] = {"pssi_sections": [
        {"title": "Politique de contrôle d'accès", "content": "MFA obligatoire pour tous les accès distants."},
    ]}
    _, contenu = report_docx.build_pssi_docx(mission, "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Politique de contrôle d'accès" in texte
    assert "MFA obligatoire pour tous les accès distants." in texte


def test_pssi_signale_l_absence_de_sections_redigees():
    _, contenu = report_docx.build_pssi_docx(_mission(), "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Aucune section PSSI n'a été rédigée à ce stade." in texte


def test_pssi_reprend_le_rto_rpo_et_la_sequence_e3r():
    mission = _mission_complete()
    _, contenu = report_docx.build_pssi_docx(mission, "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "4 h" in texte and "1 h" in texte  # RTO / RPO
    assert "Immuable" in texte  # politique de sauvegarde
    assert "Isolement" in texte  # E3R : endiguement
    assert "Priorité à l'éradication avant redémarrage" in texte  # volet stratégique (§14.2.3)


def test_pssi_porte_un_bloc_signatures_pour_homologation():
    _, contenu = report_docx.build_pssi_docx(_mission(), "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Signatures pour homologation de sécurité" in texte


# --- AIPD / PIA -------------------------------------------------------------------

def test_aipd_rend_le_registre_et_les_quatre_volets_du_pia():
    mission = _mission_complete()
    _, contenu = report_docx.build_aipd_docx(mission, "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Paie" in texte  # registre RGPD
    assert "Vidéoprotection" in texte  # volet 2.1 : description systématique


def test_aipd_signale_les_volets_du_pia_non_rediges():
    _, contenu = report_docx.build_aipd_docx(_mission(), "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "n'a pas été rédigé." in texte


def test_aipd_n_est_pas_conditionne_a_aipd_required():
    """Contrairement au signal d'alerte de la Phase 2, ce document se produit
    sur demande explicite du consultant, même si `aipd_required` est faux."""
    mission = _mission_complete()
    mission["steps"]["diagnostic"]["aipd_required"] = False
    _, contenu = report_docx.build_aipd_docx(mission, "acme")
    Document(io.BytesIO(contenu))  # ne lève pas, contenu rendu normalement


def test_aipd_porte_les_obligations_organisationnelles():
    mission = _mission_complete()
    _, contenu = report_docx.build_aipd_docx(mission, "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Avis rendu" in texte  # commentaire de l'obligation DPO


def test_aipd_porte_un_bloc_signature_cnil():
    _, contenu = report_docx.build_aipd_docx(_mission(), "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "Signature de validation conformité CNIL" in texte


# --- Robustesse transversale aux quatre livrables --------------------------------

@pytest.mark.parametrize("cle", sorted(_BATISSEURS_AUTRES))
def test_chaque_livrable_survit_aux_caracteres_speciaux(cle):
    mission = _mission_complete()
    mission["client"] = "ACME & Fils <SAS>"
    _, contenu = _BATISSEURS_AUTRES[cle](mission, "acme")
    texte = _all_text(Document(io.BytesIO(contenu)))
    assert "ACME & Fils <SAS>" in texte


@pytest.mark.parametrize("cle", sorted(_BATISSEURS_AUTRES))
def test_chaque_livrable_n_a_aucune_ligne_de_tableau_entierement_vide(cle):
    _, contenu = _BATISSEURS_AUTRES[cle](_mission_complete(), "acme", "Dorian", "DP Cyber")
    doc = Document(io.BytesIO(contenu))
    for t in doc.tables:
        for r in t.rows:
            assert any(c.text.strip() for c in r.cells), \
                f"ligne de tableau entièrement vide dans le livrable {cle!r}"
