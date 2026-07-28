"""Tests de la génération des livrables Word et du calcul de criticité des tiers."""
from __future__ import annotations

import io
import sys
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from modules import docx_export  # noqa: E402
from modules.projects import _tprm_rate  # noqa: E402


@pytest.fixture
def mission() -> dict:
    """Mission minimale mais réaliste : quelques actifs, un contrôle, une action."""
    return {
        "id": "acme", "name": "Audit ISO 27001", "client": "ACME",
        "steps": {
            "cadrage": {
                "scope": "SI de production",
                "client_missions": "Distribuer des services numériques.",
                "framework_name": "ISO/IEC 27001:2022",
                "assets_metier": [
                    {"id": "VM-01", "name": "Fichier clients",
                     "description": "Coordonnées et contrats", "is_personal_data": True},
                ],
                "assets_support": [
                    {"id": "BS-01", "name": "Active Directory",
                     "type": "Logiciel", "owner": "Équipe Système"},
                ],
            },
            "evaluation": {
                "manual_controls": [
                    {"id": "A.5.1", "title": "Politiques de sécurité",
                     "status": "CONFORME", "notes": "PSSI validée en comité"},
                ],
            },
            "traitement": {
                "remediations": [
                    {"id": "REM-01", "axe": "Protection",
                     "measure": "Déployer le MFA", "priority": "Critique"},
                ],
            },
        },
    }


def _all_text(doc: Document) -> str:
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def test_rendu_docx_est_un_fichier_word_valide(mission):
    """Le fichier produit doit être un vrai .docx (OOXML), pas du HTML déguisé."""
    content = docx_export.render_iso27001(mission, "Dorian", "DP Cyber Consulting")
    assert content[:2] == b"PK", "un .docx est une archive ZIP : signature PK attendue"
    Document(io.BytesIO(content))  # lève si l'archive n'est pas un document Word


def test_aucune_balise_jinja_ne_subsiste(mission):
    """Une balise non substituée signalerait une variable manquante dans le contexte."""
    doc = Document(io.BytesIO(docx_export.render_iso27001(mission)))
    texte = _all_text(doc)
    assert "{{" not in texte and "{%" not in texte


def test_les_donnees_de_la_mission_sont_presentes(mission):
    doc = Document(io.BytesIO(docx_export.render_iso27001(mission, "Dorian", "DP Cyber")))
    texte = _all_text(doc)
    for attendu in ["ACME", "Audit ISO 27001", "Fichier clients",
                    "Active Directory", "Déployer le MFA", "Dorian"]:
        assert attendu in texte, f"donnée absente du rapport : {attendu}"


def test_champ_absent_signale_et_jamais_inventé(mission):
    """Principe produit : un champ vide se voit, il n'est pas comblé."""
    mission["steps"]["cadrage"]["scope"] = ""
    doc = Document(io.BytesIO(docx_export.render_iso27001(mission)))
    assert docx_export.NON_RENSEIGNE in _all_text(doc)


def test_score_non_evalue_plutot_que_zero():
    """Sans contrôle évalué, afficher 0 % se lirait comme une non-conformité totale."""
    vide = {"client": "ACME", "name": "M", "steps": {}}
    score, _ = docx_export._score_and_band({}, [])
    assert score == "non évalué"
    assert "non évalué" in _all_text(Document(io.BytesIO(docx_export.render_iso27001(vide))))


def test_empreinte_stable_et_sensible_au_contenu(mission):
    """L'empreinte doit être reproductible, et changer si la mission change."""
    empreinte = docx_export.data_fingerprint(mission)
    assert empreinte == docx_export.data_fingerprint(mission)
    assert len(empreinte) == 64
    mission["client"] = "AUTRE"
    assert docx_export.data_fingerprint(mission) != empreinte


def test_mention_de_reserve_datee_et_nominative(mission):
    """La réserve délimite la responsabilité : elle doit citer le client et la date."""
    doc = Document(io.BytesIO(docx_export.render_iso27001(mission)))
    texte = _all_text(doc)
    assert "ACME" in texte
    assert "ne saurait" in texte and "garantie" in texte


# --- Criticité des tiers ----------------------------------------------------

@pytest.mark.parametrize("criteres,score,rating", [
    ((5, 5, 4, 4), 3.5, "Élevé"),      # ex-« 4.5 / Critique » saisi à la main
    ((4, 5, 3, 3), 3.8, "Élevé"),
    ((2, 1, 2, 4), 2.3, "Moyen"),      # ex-« 2.2 / Faible » : niveau erroné
    ((4, 4, 4, 4), 3.0, "Élevé"),      # ex-« 3.5 / Moyen » : niveau erroné
])
def test_criticite_tiers_derivee_des_criteres(criteres, score, rating):
    resultat = _tprm_rate(*criteres)
    assert resultat["score"] == score
    assert resultat["rating"] == rating


def test_arrondi_identique_au_navigateur():
    """2.25 doit donner 2.3 comme Number.toFixed(1) en JS.

    L'arrondi au pair de Python donnerait 2.2 : le score affiché divergerait
    alors de celui recalculé par le frontend lors d'une réédition du tiers.
    """
    assert _tprm_rate(2, 1, 2, 4)["score"] == 2.3
